"""
Сервис для парсинга документов
Извлекает текст и метаданные из PDF, DOCX, Excel, Markdown
"""

import os
import logging
from typing import Dict, Any, Optional, List
from pathlib import Path
import tempfile

# PDF парсинг
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

# DOCX парсинг
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Excel парсинг
try:
    import openpyxl
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False

# Markdown парсинг
try:
    import markdown
    HAS_MARKDOWN = True
except ImportError:
    HAS_MARKDOWN = False

logger = logging.getLogger(__name__)


class DocumentParser:
    """Парсер для различных форматов документов"""
    
    def __init__(self):
        """Инициализация парсера"""
        self.supported_formats = {
            'pdf': HAS_PDF,
            'docx': HAS_DOCX,
            'xlsx': HAS_EXCEL,
            'md': HAS_MARKDOWN,
            'txt': True  # Всегда поддерживается
        }
        
        logger.info(f"DocumentParser initialized. Supported formats: {self.supported_formats}")
    
    async def parse_file(self, file_path: str, file_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Универсальный метод для парсинга файла
        
        Args:
            file_path: Путь к файлу
            file_type: Тип файла (pdf, docx, xlsx, md, txt)
        
        Returns:
            Dict с извлеченными данными
        """
        # Определяем тип файла по расширению если не указан
        if not file_type:
            file_type = Path(file_path).suffix.lower().lstrip('.')
        
        # Выбираем подходящий парсер
        parsers = {
            'pdf': self.parse_pdf,
            'docx': self.parse_docx,
            'doc': self.parse_docx,
            'xlsx': self.parse_excel,
            'xls': self.parse_excel,
            'md': self.parse_markdown,
            'markdown': self.parse_markdown,
            'txt': self.parse_text
        }
        
        parser = parsers.get(file_type)
        
        if not parser:
            return {
                "success": False,
                "error": f"Unsupported file type: {file_type}",
                "supported_types": list(parsers.keys())
            }
        
        try:
            result = await parser(file_path)
            result['file_type'] = file_type
            result['success'] = True
            return result
        except Exception as e:
            logger.error(f"Failed to parse {file_type} file: {e}")
            return {
                "success": False,
                "error": str(e),
                "file_type": file_type
            }
    
    async def parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Парсит PDF файл
        
        Returns:
            Dict с текстом, метаданными и информацией о страницах
        """
        if not HAS_PDF:
            raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Извлекаем метаданные
            metadata = pdf_reader.metadata or {}
            
            # Извлекаем текст со всех страниц
            pages_text = []
            total_text = []
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                pages_text.append({
                    "page_number": page_num + 1,
                    "text": page_text,
                    "char_count": len(page_text)
                })
                total_text.append(page_text)
            
            full_text = "\n\n".join(total_text)
            
            return {
                "text": full_text,
                "pages": pages_text,
                "page_count": len(pdf_reader.pages),
                "metadata": {
                    "title": metadata.get('/Title', ''),
                    "author": metadata.get('/Author', ''),
                    "subject": metadata.get('/Subject', ''),
                    "creator": metadata.get('/Creator', ''),
                    "created_at": metadata.get('/CreationDate', '')
                },
                "char_count": len(full_text),
                "word_count": len(full_text.split())
            }
    
    async def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Парсит DOCX файл
        
        Returns:
            Dict с текстом, параграфами и метаданными
        """
        if not HAS_DOCX:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        doc = Document(file_path)
        
        # Извлекаем параграфы
        paragraphs = []
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal"
                })
                full_text.append(para.text)
        
        # Извлекаем таблицы
        tables_data = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                table_rows.append([cell.text for cell in row.cells])
            tables_data.append(table_rows)
        
        text = "\n\n".join(full_text)
        
        return {
            "text": text,
            "paragraphs": paragraphs,
            "paragraph_count": len(paragraphs),
            "tables": tables_data,
            "table_count": len(tables_data),
            "char_count": len(text),
            "word_count": len(text.split()),
            "metadata": {
                "has_tables": len(tables_data) > 0,
                "has_images": len(doc.inline_shapes) > 0
            }
        }
    
    async def parse_excel(self, file_path: str) -> Dict[str, Any]:
        """
        Парсит Excel файл
        
        Returns:
            Dict с данными из всех листов
        """
        if not HAS_EXCEL:
            raise ImportError("openpyxl not installed. Install with: pip install openpyxl")
        
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        
        sheets_data = []
        all_text = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            # Извлекаем данные
            rows_data = []
            for row in sheet.iter_rows(values_only=True):
                # Пропускаем пустые строки
                if any(cell is not None for cell in row):
                    row_text = [str(cell) if cell is not None else '' for cell in row]
                    rows_data.append(row_text)
                    all_text.append(" | ".join(row_text))
            
            sheets_data.append({
                "sheet_name": sheet_name,
                "rows": rows_data,
                "row_count": len(rows_data),
                "column_count": sheet.max_column
            })
        
        text = "\n".join(all_text)
        
        return {
            "text": text,
            "sheets": sheets_data,
            "sheet_count": len(sheets_data),
            "total_rows": sum(s['row_count'] for s in sheets_data),
            "char_count": len(text),
            "metadata": {
                "workbook_name": Path(file_path).stem
            }
        }
    
    async def parse_markdown(self, file_path: str) -> Dict[str, Any]:
        """
        Парсит Markdown файл
        
        Returns:
            Dict с текстом, HTML и метаданными
        """
        with open(file_path, 'r', encoding='utf-8') as file:
            markdown_text = file.read()
        
        # Конвертируем в HTML если библиотека доступна
        html_content = ""
        if HAS_MARKDOWN:
            html_content = markdown.markdown(markdown_text)
        
        # Извлекаем заголовки
        headers = []
        for line in markdown_text.split('\n'):
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                headers.append({
                    "level": level,
                    "text": text
                })
        
        return {
            "text": markdown_text,
            "html": html_content,
            "headers": headers,
            "header_count": len(headers),
            "char_count": len(markdown_text),
            "word_count": len(markdown_text.split()),
            "metadata": {
                "has_code_blocks": "```" in markdown_text,
                "has_links": "[" in markdown_text and "](" in markdown_text
            }
        }
    
    async def parse_text(self, file_path: str) -> Dict[str, Any]:
        """
        Парсит обычный текстовый файл
        
        Returns:
            Dict с текстом и базовой статистикой
        """
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        lines = text.split('\n')
        
        return {
            "text": text,
            "lines": lines,
            "line_count": len(lines),
            "char_count": len(text),
            "word_count": len(text.split()),
            "metadata": {
                "encoding": "utf-8"
            }
        }
    
    async def extract_key_points(self, parsed_data: Dict[str, Any], max_points: int = 5) -> List[str]:
        """
        Извлекает ключевые пункты из распарсенного документа
        
        Args:
            parsed_data: Результат парсинга
            max_points: Максимальное количество пунктов
        
        Returns:
            Список ключевых пунктов
        """
        text = parsed_data.get('text', '')
        
        # Простое извлечение - берем первые предложения
        sentences = text.split('.')
        key_points = []
        
        for sentence in sentences[:max_points]:
            sentence = sentence.strip()
            if len(sentence) > 20:  # Пропускаем слишком короткие
                key_points.append(sentence)
        
        return key_points


# Singleton instance
_document_parser = None


def get_document_parser() -> DocumentParser:
    """Получить singleton instance Document Parser"""
    global _document_parser
    if _document_parser is None:
        _document_parser = DocumentParser()
    return _document_parser

